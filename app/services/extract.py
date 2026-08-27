"""Text extraction from uploaded evidence files.

Supports PDF (PyMuPDF, with OCR fallback for scans via pytesseract when
available), DOCX (python-docx) and plain text. Extraction is best-effort:
an unreadable file yields empty text and the evidence agent responds with
an explicit "no text could be extracted" note instead of guessing.
"""
from __future__ import annotations

import asyncio
import io
import logging
from pathlib import Path

logger = logging.getLogger("nyaya.extract")

try:  # PyMuPDF (modern import name; `fitz` kept as legacy fallback)
    import pymupdf  # type: ignore

    PDF_AVAILABLE = True
except Exception:  # pragma: no cover
    try:
        import fitz as pymupdf  # type: ignore

        PDF_AVAILABLE = True
    except Exception:
        pymupdf = None  # type: ignore
        PDF_AVAILABLE = False

try:  # python-docx
    import docx  # type: ignore

    DOCX_AVAILABLE = True
except Exception:  # pragma: no cover
    docx = None  # type: ignore
    DOCX_AVAILABLE = False

try:  # optional OCR for scanned PDFs
    import pytesseract  # type: ignore
    from PIL import Image  # type: ignore

    OCR_AVAILABLE = True
except Exception:  # pragma: no cover
    pytesseract = None  # type: ignore
    OCR_AVAILABLE = False

TEXT_SUFFIXES = {".txt", ".md", ".csv", ".log"}


def _pdf_text(path: str) -> str:
    if not PDF_AVAILABLE:
        logger.warning("PyMuPDF not installed — cannot read PDF: %s", path)
        return ""
    parts: list[str] = []
    with pymupdf.open(path) as pdf:
        for page in pdf:
            parts.append(page.get_text())
    text = "\n".join(parts).strip()
    if text:
        return text
    # No embedded text -> likely a scan. OCR if tesseract is available.
    if not OCR_AVAILABLE:
        logger.info("PDF has no embedded text and OCR is unavailable: %s", path)
        return ""
    ocr_parts: list[str] = []
    with pymupdf.open(path) as pdf:
        for page in pdf:
            pix = page.get_pixmap(dpi=200)
            image = Image.open(io.BytesIO(pix.tobytes("ppm")))
            ocr_parts.append(pytesseract.image_to_string(image))
    return "\n".join(ocr_parts).strip()


def _docx_text(path: str) -> str:
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not installed — cannot read DOCX: %s", path)
        return ""
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


async def extract_text(path: str, mime: str = "") -> str:
    """Extract text from a stored upload. Returns "" when nothing is readable."""
    suffix = Path(path).suffix.lower()
    try:
        if suffix == ".pdf":
            return await asyncio.to_thread(_pdf_text, path)
        if suffix == ".docx":
            return await asyncio.to_thread(_docx_text, path)
        if suffix in TEXT_SUFFIXES:
            return Path(path).read_text(encoding="utf-8", errors="ignore").strip()
        if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
            if OCR_AVAILABLE:
                def _ocr() -> str:
                    return pytesseract.image_to_string(Image.open(path))
                return await asyncio.to_thread(_ocr)
            return ""
        logger.info("Unsupported evidence format (%s): %s", suffix, path)
        return ""
    except Exception:
        logger.exception("Text extraction failed for %s", path)
        return ""
